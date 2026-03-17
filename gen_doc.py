from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 页面边距 ──────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3.17)
section.right_margin  = Cm(3.17)
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)

# ── 样式辅助函数 ──────────────────────────────────────────
def set_font(run, name='宋体', size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), 'Times New Roman' if name == '宋体' else name)
    rPr.insert(0, rFonts)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    sizes = {1: 16, 2: 14, 3: 13}
    set_font(run, name='黑体', size=sizes.get(level, 12), bold=True)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    if level == 1:
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '4472C4')
        pBdr.append(bottom)
        pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, name='宋体', size=12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(20)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    return p

def add_formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, name='Times New Roman', size=11)
    run.font.italic = True
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    hrow = table.rows[0]
    for j, h in enumerate(headers):
        cell = hrow.cells[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, name='黑体', size=11, bold=True)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'BDD7EE')
        tcPr.append(shd)
    # 数据行
    for i, row in enumerate(rows):
        tr = table.rows[i + 1]
        for j, val in enumerate(row):
            cell = tr.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, name='宋体', size=11)
            if i % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F2F7FC')
                tcPr.append(shd)
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Cm(w)
    return table

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, name='宋体', size=12)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(20)
    return p

# ══════════════════════════════════════════════════════════
#  封面
# ══════════════════════════════════════════════════════════
doc.add_paragraph('\n\n')
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('涡扇发动机效能参数贝叶斯辨识报告')
set_font(run, name='黑体', size=22, bold=True, color=(31, 78, 121))

doc.add_paragraph('\n')
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run('——基于自适应 Metropolis-Hastings MCMC 方法')
set_font(run, name='宋体', size=14, color=(89, 89, 89))

doc.add_paragraph('\n\n\n')
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_p.add_run('日期：2026 年 3 月')
set_font(run, name='宋体', size=12, color=(89, 89, 89))
doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  一、问题描述
# ══════════════════════════════════════════════════════════
add_heading(doc, '一、问题描述', 1)

add_heading(doc, '1.1 工程背景', 2)
add_body(doc,
    '涡扇发动机的性能由一组内部热力学参数共同决定，这些参数（如各部件效率、'
    '总压恢复系数）在实际工程中难以直接测量，只能通过外部可观测的整机性能指标'
    '（单位比推力 R_ud 和单位比油耗 C_ud）进行间接辨识。由于测量噪声的存在，'
    '参数辨识问题具有不确定性，因此需要借助贝叶斯推断框架对参数的后验分布进行定量估计。',
    indent=True)

add_heading(doc, '1.2 待辨识参数', 2)
add_body(doc, '共涉及 13 个参数，均为无量纲效率或总压恢复系数，其先验范围（均匀分布上下界）如下表所示：', indent=True)

param_headers = ['编号', '参数符号', '物理含义', '下界 lb', '上界 ub', '真值 θ*']
param_rows = [
    [1,  'η_k',       '压气机绝热效率',       '0.84',  '0.86',  '0.8500'],
    [2,  'η_t',       '涡轮绝热效率',         '0.86',  '0.92',  '0.8900'],
    [3,  'η_m',       '机械效率',             '0.980', '0.995', '0.9880'],
    [4,  'η_v',       '风扇效率',             '0.85',  '0.87',  '0.8600'],
    [5,  'η_tv',      '风扇涡轮效率',         '0.90',  '0.92',  '0.9100'],
    [6,  'η_c1',      '一次喷管效率',         '0.94',  '0.95',  '0.9450'],
    [7,  'η_c2',      '二次喷管效率',         '0.92',  '0.94',  '0.9300'],
    [8,  'σ_cc',      '进气道激波总压恢复系数', '0.98',  '1.00',  '0.9900'],
    [9,  'σ_kan',     '进气通道总压恢复系数',  '0.98',  '0.99',  '0.9850'],
    [10, 'σ_kask',    '压气机级间总压恢复系数','0.98',  '0.99',  '0.9850'],
    [11, 'σ_ks',      '燃烧室总压恢复系数',   '0.94',  '0.96',  '0.9500'],
    [12, 'η_T',       '燃烧放热系数',         '0.97',  '0.99',  '0.9800'],
    [13, 'λ',         '风扇涡轮热恢复系数',   '1.02',  '1.04',  '1.0300'],
]
add_table(doc, param_headers, param_rows, col_widths=[1.2, 1.6, 4.5, 1.4, 1.4, 1.6])
doc.add_paragraph()

add_heading(doc, '1.3 工作条件', 2)
cond_headers = ['参数', '符号', '取值']
cond_rows = [
    ['大气温度',     'T_H',       '288 K'],
    ['飞行马赫数',   'M_∞',       '0.0（地面静止）'],
    ['涵道比',       'm',         '10.0'],
    ['压气机总压比', 'π_k',       '33.0'],
    ['涡轮前燃气温度','T_g',      '1700 K'],
]
add_table(doc, cond_headers, cond_rows, col_widths=[4.0, 3.0, 5.2])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  二、前向模型
# ══════════════════════════════════════════════════════════
add_heading(doc, '二、前向模型', 1)
add_body(doc,
    '前向模型 engine_forward(θ, cond) 描述从参数向量 θ 到可观测性能指标的映射：',
    indent=True)
add_formula(doc, 'y = F(θ) = [R_ud,  C_ud]ᵀ')

add_heading(doc, '2.1 主要计算步骤', 2)

steps = [
    ('步骤 1：飞行速度与进口驻点温度',
     'a = √(k_air · R_air · T_H),   V_∞ = a · M_∞\n'
     'τ_v = (1 + V_∞² / (2·(k_air/(k_air-1))·R_air·T_H))^(k_air/(k_air-1))\n'
     'T_B = T_H · τ_v'),
    ('步骤 2：压气机出口温度',
     'T_k = T_B · (1 + (π_k^((k_air-1)/k_air) - 1) / η_k)'),
    ('步骤 3：相对耗油量',
     'g_T = 3×10⁻⁵ · T_g  −  2.69×10⁻⁵ · T_k  −  0.003'),
    ('步骤 4：热恢复系数',
     'W_c = k_air/(k_air-1) · R_air · T_B · (π_k^((k_air-1)/k_air) − 1)\n'
     'h_g = k_T/(k_T-1) · R_T · T_g\n'
     'λ_heat = (1 − W_c/(h_g·η_k))  /  (1 − W_c/(h_g·η_k·η_t))'),
    ('步骤 5：单位自由能',
     'L_sv = λ_heat · [ k_T/(k_T-1)·R_T·T_g·(1−(τ_v·σ_bx·π_k·σ_kask·σ_ks)^(-(k_T-1)/k_T))\n'
     '         −  W_c / ((1+g_T)·η_k·η_T·η_t·η_m·(1−δ)) ]'),
    ('步骤 6：最优能量分配系数',
     'x_pc = (1 + m·V_∞²/(2·L_sv·η_tv·η_v·η_c2))\n'
     '       / (1 + m·η_tv·η_v·η_c2 / (η_c1·λ))'),
    ('步骤 7：单位比推力与比油耗',
     'V_j1 = (1+g_T)·√(2·η_c1·λ·x_pc·L_sv) − V_∞\n'
     'V_j2 = √(2·(1−x_pc)/m·L_sv·η_tv·η_v·η_c2 + V_∞²) − V_∞\n'
     'R_ud = V_j1/(1+m)  +  m·V_j2/(1+m)\n'
     'C_ud = 3600·g_T·(1−δ) / (R_ud·(1+m))'),
]
for title, formula in steps:
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_font(run, name='黑体', size=12, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    add_formula(doc, formula)

add_heading(doc, '2.2 分段物性参数', 2)
add_body(doc, '燃气绝热指数 k_T 和气体常数 R_T 随温度按以下规则取值：', indent=True)
phys_headers = ['T_g 范围', 'k_T', 'R_T (J/(kg·K))']
phys_rows = [
    ['≤ 1400 K', '1.33', '287.6'],
    ['1400–1600 K', '1.30', '288.0'],
    ['> 1600 K',  '1.25', '288.6'],
]
add_table(doc, phys_headers, phys_rows, col_widths=[5.0, 3.0, 4.2])
doc.add_paragraph()
add_body(doc, '冷却引气系数：δ = clip(0.02 + (T_g − 1200)/100 × 0.02,  0,  0.15)', indent=True)

# ══════════════════════════════════════════════════════════
#  三、算法设置
# ══════════════════════════════════════════════════════════
add_heading(doc, '三、算法设置', 1)

add_heading(doc, '3.1 贝叶斯框架', 2)

add_heading(doc, '3.1.1 先验分布', 3)
add_body(doc, '对参数向量 θ 采用均匀先验，反映工程专家对参数物理范围的约束：', indent=True)
add_formula(doc, 'ln p(θ) = 0   if θ ∈ [lb, ub]；  −∞  otherwise')

add_heading(doc, '3.1.2 似然函数', 3)
add_body(doc, '观测模型假设加性高斯噪声，噪声标准差取真值的 1%：', indent=True)
add_formula(doc, 'R_obs = R_ud(θ) + ε_R,   ε_R ~ N(0, σ_R²),   σ_R = 0.01·|R_true|')
add_formula(doc, 'C_obs = C_ud(θ) + ε_C,   ε_C ~ N(0, σ_C²),   σ_C = 0.01·|C_true|')
add_body(doc, '对数似然：', indent=True)
add_formula(doc,
    'ln p(y_obs | θ) = −0.5·[(R_obs−R_ud)²/σ_R² + ln(2πσ_R²)\n'
    '                        + (C_obs−C_ud)²/σ_C² + ln(2πσ_C²)]')

add_heading(doc, '3.1.3 后验分布', 3)
add_formula(doc, 'ln p(θ | y_obs) = ln p(θ) + ln p(y_obs | θ)')

add_heading(doc, '3.2 MCMC 算法：自适应 Metropolis-Hastings', 2)

add_heading(doc, '3.2.1 参数空间变换', 3)
add_body(doc,
    '为天然满足参数有界约束，将原始参数映射到单位超立方体，在 z 空间上执行随机游走，'
    '越界时采用反射边界处理：',
    indent=True)
add_formula(doc, 'z_i = (θ_i − lb_i) / (ub_i − lb_i) ∈ [0, 1]')
add_formula(doc, '若 z < 0 → z = −z；  若 z > 1 → z = 2 − z')

add_heading(doc, '3.2.2 提议分布与接受准则', 3)
add_formula(doc, 'z* = z^(t) + ε,   ε ~ N(0, σ_prop²·I)')
add_formula(doc, '接受条件：ln u < ln p(θ*|y) − ln p(θ^(t)|y),   u ~ U[0,1]')

add_heading(doc, '3.2.3 自适应步长规则', 3)
add_body(doc,
    '每隔 100 步统计窗口内接受率，若偏低（< 0.18）则缩减步长 ×0.9，'
    '若偏高（> 0.35）则放大步长 ×1.1，步长限幅在 [1×10⁻⁴, 0.5] 内。'
    '自适应仅在烧入期（步数 500–2500）内进行，后验采样阶段步长固定。',
    indent=True)

add_heading(doc, '3.2.4 采样配置汇总', 3)
cfg_headers = ['配置项', '取值']
cfg_rows = [
    ['总采样步数',     '10,000'],
    ['烧入期（Burn-in）', '2,000 步'],
    ['后验有效链长度', '8,000 步'],
    ['初始提议标准差 σ_prop', '0.012'],
    ['自适应区间',     '步数 500–2500'],
    ['调整间隔',       '每 100 步'],
    ['目标接受率区间', '[0.18, 0.35]'],
    ['步长调整因子',   '偏低 → ×0.9；偏高 → ×1.1'],
    ['σ_prop 限幅',    '[1×10⁻⁴, 0.5]'],
    ['初始点 θ₀',      '各参数先验区间中点'],
    ['随机种子',       '42（主程序）/ 123（数据生成）'],
]
add_table(doc, cfg_headers, cfg_rows, col_widths=[6.0, 6.2])
doc.add_paragraph()

add_heading(doc, '3.3 后验统计量', 2)
add_body(doc, '从后验链 chain_post（去除烧入期后的 8000 步）提取以下统计量：', indent=True)
for item in [
    '后验均值：θ_mean = (1/N)·Σ θ^(s)',
    'MAP 估计：后验链中对数后验最大处的样本  θ_MAP = argmax_s ln p(θ^(s)|y)',
    '95% 可信区间：[quantile(2.5%), quantile(97.5%)]',
]:
    add_bullet(doc, item)

# ══════════════════════════════════════════════════════════
#  四、结果
# ══════════════════════════════════════════════════════════
add_heading(doc, '四、结果', 1)

add_heading(doc, '4.1 前向模型真值输出', 2)
add_body(doc,
    '在真值参数 θ* 和给定工况下（地面静止，M_∞ = 0），前向模型的关键中间量为：',
    indent=True)
fwd_headers = ['中间量', '符号', '说明']
fwd_rows = [
    ['进口驻点温度',   'T_B',       '= T_H = 288 K（静止时 τ_v = 1）'],
    ['压气机出口温度', 'T_k',       '由 π_k = 33、η_k = 0.85 计算，约 800–900 K'],
    ['冷却引气系数',   'δ',         '= 0.02+(1700−1200)/100×0.02 = 0.12'],
    ['相对耗油量',     'g_T',       '满足 g_T > 0（物理约束）'],
    ['热恢复系数',     'λ_heat',    '由压气机功与燃气焓之比决定'],
    ['单位自由能',     'L_sv',      '涡轮膨胀做功与压气机功之差经热恢复修正'],
]
add_table(doc, fwd_headers, fwd_rows, col_widths=[3.5, 2.0, 6.7])
doc.add_paragraph()

add_heading(doc, '4.2 后验估计结果汇总', 2)
add_body(doc,
    '下表给出各参数的贝叶斯估计结果（代码标准输出格式，实际数值由 MATLAB 运行结果填入）：',
    indent=True)
post_headers = ['参数', '真值', '后验均值', 'MAP', 'CI95 低', 'CI95 高', '先验中点']
post_rows = [
    ['η_k',    '0.8500', '—', '—', '—', '—', '0.8500'],
    ['η_t',    '0.8900', '—', '—', '—', '—', '0.8900'],
    ['η_m',    '0.9880', '—', '—', '—', '—', '0.9875'],
    ['η_v',    '0.8600', '—', '—', '—', '—', '0.8600'],
    ['η_tv',   '0.9100', '—', '—', '—', '—', '0.9100'],
    ['η_c1',   '0.9450', '—', '—', '—', '—', '0.9450'],
    ['η_c2',   '0.9300', '—', '—', '—', '—', '0.9300'],
    ['σ_cc',   '0.9900', '—', '—', '—', '—', '0.9900'],
    ['σ_kan',  '0.9850', '—', '—', '—', '—', '0.9850'],
    ['σ_kask', '0.9850', '—', '—', '—', '—', '0.9850'],
    ['σ_ks',   '0.9500', '—', '—', '—', '—', '0.9500'],
    ['η_T',    '0.9800', '—', '—', '—', '—', '0.9800'],
    ['λ',      '1.0300', '—', '—', '—', '—', '1.0300'],
]
add_table(doc, post_headers, post_rows, col_widths=[2.0, 1.6, 2.0, 1.8, 1.8, 1.8, 1.8])
doc.add_paragraph()
add_body(doc, '注：表中"—"处由 MATLAB 运行后实际输出填入。', indent=True)

add_heading(doc, '4.3 后验预测对比', 2)
pred_headers = ['性能指标', '真值', '观测值（含1%噪声）', '后验均值预测', 'MAP 预测']
pred_rows = [
    ['R_ud (N·s/kg)',    'R*', 'R* + ε_R', 'F₁(θ_mean)', 'F₁(θ_MAP)'],
    ['C_ud (kg/(N·h))', 'C*', 'C* + ε_C', 'F₂(θ_mean)', 'F₂(θ_MAP)'],
]
add_table(doc, pred_headers, pred_rows, col_widths=[3.5, 2.0, 3.5, 3.0, 2.8])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  五、分析
# ══════════════════════════════════════════════════════════
add_heading(doc, '五、分析', 1)

add_heading(doc, '5.1 算法收敛性', 2)
add_body(doc,
    '链轨迹（Trace Plot）是判断 MCMC 收敛的主要可视化工具。收敛良好的链应表现为：'
    '烧入期后在参数真值附近平稳振荡，无明显趋势漂移；链的探索范围覆盖后验分布的'
    '主要质量区域；自相关长度适中。',
    indent=True)
add_body(doc,
    '接受率是评估提议分布质量的关键指标。理论上，对于高维问题最优接受率约为 23%，'
    '本代码目标区间 [0.18, 0.35] 与此一致。',
    indent=True)

add_heading(doc, '5.2 可辨识性与参数敏感性', 2)
add_body(doc,
    '由于输出观测量只有两个（R_ud、C_ud），而待辨识参数多达 13 个，系统存在欠定性：',
    indent=True)
for item in [
    '强可辨识参数（η_k、η_t、η_T 等）：对输出有显著单调影响，后验分布较集中，'
    '95% 可信区间明显窄于先验区间。',
    '弱可辨识参数（部分总压恢复系数）：对输出影响较小，后验分布接近均匀先验，'
    '说明数据对这些参数的约束力弱。',
    '参数相关性：后验分布中可能出现参数间的正/负相关，意味着数据只能约束参数的某种组合。',
]:
    add_bullet(doc, item)

add_heading(doc, '5.3 噪声水平的影响', 2)
add_body(doc,
    '本实验噪声水平为 1%，属于工程中较低噪声。噪声越小，似然函数越尖锐，后验越集中，'
    '辨识精度越高。若噪声增大至 5%–10%，则后验分布变宽，弱可辨识参数的后验更趋近先验。',
    indent=True)

add_heading(doc, '5.4 均匀先验的合理性', 2)
add_body(doc,
    '均匀先验反映工程专家知识：参数被限制在合理物理范围内，但范围内无偏好，'
    '是一种保守而客观的假设。若有更精确的先验信息（如高斯先验），'
    '可进一步收窄后验不确定性。',
    indent=True)

add_heading(doc, '5.5 后验均值 vs. MAP 估计', 2)
add_body(doc,
    '后验均值 θ_mean 是 L₂ 损失下的贝叶斯最优估计，对后验形状整体敏感，'
    '适合后验接近对称分布的情形。MAP 估计 θ_MAP 是后验的众数，'
    '链长越短越不稳定。当后验接近单峰对称高斯时，两者趋于一致；'
    '差异较大则说明后验存在明显偏斜或多峰结构。',
    indent=True)

add_heading(doc, '5.6 反射边界的作用', 2)
add_body(doc,
    '参数有界约束若用截断或拒绝方案处理，会导致边界处采样不足。'
    '本代码采用反射边界，在 [0,1] 范围内对 z 空间作折叠反射，'
    '保证了边界附近的正确采样密度，是处理有界均匀先验的标准方法。',
    indent=True)

add_heading(doc, '5.7 局限性与改进方向', 2)
lim_headers = ['局限性', '潜在改进方向']
lim_rows = [
    ['单链 MCMC，难以诊断收敛',       '多链并行，计算 Gelman-Rubin 统计量 R̂'],
    ['各维度共享步长，适应性有限',    '全协方差自适应 MCMC（AM-MCMC）'],
    ['仅有两个观测量，系统欠定',      '增加多工况（不同 M_∞、π_k、T_g）数据'],
    ['MAP 由链内最优样本确定，精度受链长限制', '后处理阶段用梯度优化精化 MAP'],
    ['链相关性未量化',                '计算有效样本量（ESS）与积分自相关时间（IACT）'],
]
add_table(doc, lim_headers, lim_rows, col_widths=[6.3, 6.1])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  六、结论
# ══════════════════════════════════════════════════════════
add_heading(doc, '六、结论', 1)
add_body(doc,
    '本工作构建了涡扇发动机热力学参数的贝叶斯辨识框架，核心贡献包括：',
    indent=True)
for item in [
    '前向模型：实现了从 13 个效率/总压恢复参数到单位比推力和比油耗的完整热力学计算链，'
    '包含分段物性参数和涡轮冷却引气效应。',
    '贝叶斯推断：在均匀先验（参数物理范围约束）与高斯似然（1% 测量噪声）框架下，'
    '建立了参数的后验分布模型。',
    '自适应 MCMC：采用带反射边界的自适应 Metropolis-Hastings 算法，'
    '在变换空间内高效采样 10,000 步（有效后验链 8,000 步），'
    '自适应步长确保接受率在合理范围内。',
    '不确定性量化：提供后验均值、MAP 估计及 95% 可信区间，'
    '量化了单次测量条件下各参数的辨识不确定性。',
]:
    add_bullet(doc, item)

add_body(doc,
    '该框架为航空发动机健康监测与性能退化诊断提供了方法论基础，'
    '可扩展至多工况数据融合与实时参数跟踪场景。',
    indent=True)

# ── 保存 ───────────────────────────────────────────────────
out_path = '/home/user/BiShe666/涡扇发动机参数辨识_MCMC报告.docx'
doc.save(out_path)
print(f'Done: {out_path}')
